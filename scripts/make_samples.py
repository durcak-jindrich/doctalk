"""Regenerate the binary sample documents in `samples/`.

The DOCX and PDF demo files are committed, but binaries are unreviewable in a
diff — so their text lives here and the files are generated from it. Editing a
sample means editing this script and re-running it.

    uv run --with reportlab python -m scripts.make_samples

`reportlab` is only needed to author the PDF and is deliberately not a project
dependency: nothing at runtime writes PDFs. `python-docx` is already a
dependency (the DOCX *parser* needs it).

Content is chosen to exercise both parser paths for real:

* The DOCX carries `Heading 1`/`Heading 2` styles, which `app/parsers/docx.py`
  turns into `section_path` — and the chunker starts a new chunk at every
  section boundary, so each heading yields at least one citable chunk.
* The PDF has no extractable heading structure (`section_path` is always
  `None` there), so it relies on length instead: enough prose, over enough
  pages, to pack several chunks and to show `page_number` varying in a
  citation.

Neither document mentions parental leave, vacation, passwords or refunds —
those belong to the Markdown samples, and the demo script's refusal case
("What is the parental leave allowance?") only stays honest while nothing in
the workspace answers it.
"""

from pathlib import Path

from docx import Document

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "samples"

#: (heading, paragraphs) — rendered as Heading 2 sections under one title.
ONBOARDING: list[tuple[str, list[str]]] = [
    (
        "Before Day One",
        [
            "The hiring manager raises an onboarding ticket at least five working days "
            "before a new joiner's start date. The ticket triggers account creation, "
            "equipment allocation and building access in a single workflow, so nothing "
            "depends on a manager remembering to request it separately.",
            "Every new joiner is assigned an onboarding buddy from the same team but a "
            "different project. The buddy is the default answer to a question that feels "
            "too small to ask a manager, and holds a thirty-minute check-in in each of "
            "the first four weeks.",
        ],
    ),
    (
        "First Week",
        [
            "Day one is deliberately light: a welcome session, a team lunch and the "
            "security induction. Project work starts on day two, after accounts and "
            "equipment are confirmed working.",
            "By the end of the first week a new joiner is expected to have completed the "
            "mandatory security induction, met every member of their immediate team, and "
            "shipped one small change end to end. The last item matters more than its "
            "size: it proves the toolchain works before anything urgent depends on it.",
        ],
    ),
    (
        "Equipment",
        [
            "Standard issue is a company laptop, an external monitor and a headset. IT "
            "ships equipment to arrive one working day before the start date for remote "
            "joiners, or has it waiting at the desk for office-based joiners.",
            "Requests for non-standard equipment go through the hiring manager and are "
            "approved on the basis of the role, not preference. Accessibility-related "
            "requests are exempt from that approval and are fulfilled directly by IT.",
        ],
    ),
    (
        "Expenses",
        [
            "Expense claims are submitted through the finance portal by the fifth working "
            "day of the month following the spend. Claims submitted later are paid in the "
            "next cycle rather than rejected.",
            "Receipts are required for every claim over ten euros. Travel booked through "
            "the company travel desk is billed centrally and must not be claimed as an "
            "expense.",
        ],
    ),
    (
        "Probation Review",
        [
            "The probation period is six months and ends with a documented review against "
            "the objectives agreed in the first month. There are no surprises by design: "
            "a concern raised at the six-month review should already have been raised in a "
            "regular one-to-one.",
            "A probation period may be extended once, by up to three months, where the "
            "objectives were disrupted by circumstances outside the new joiner's control.",
        ],
    ),
]

#: (heading, paragraphs). The PDF's headings are visual only — the parser sees
#: flat text, so chunking here is driven by length, not structure.
RETENTION: list[tuple[str, list[str]]] = [
    (
        "Purpose and Scope",
        [
            "This policy states how long the company keeps the data it collects, who is "
            "accountable for deleting it, and what happens when a retention period ends. "
            "It applies to every production system operated by the engineering "
            "organisation, including databases, object storage, message queues and log "
            "aggregation, whether hosted in a public cloud or on company hardware.",
            "It does not apply to material held on personal devices, to documents held in "
            "the corporate document management system, or to source code repositories. "
            "Those are covered by separate policies owned by the IT and legal functions "
            "respectively. Where two policies appear to conflict, the stricter retention "
            "period applies until the conflict is resolved by the data protection officer.",
        ],
    ),
    (
        "Retention Schedule",
        [
            "Customer account records are retained for seven years after the account is "
            "closed. The period is set by financial reporting obligations rather than by "
            "operational need, and it cannot be shortened by a team acting on its own "
            "judgement. Records within the retention window are moved to cold storage "
            "twelve months after closure, which reduces cost without reducing "
            "availability for an audit.",
            "Application logs are retained for sixty days in searchable storage and are "
            "then deleted outright. Sixty days is a deliberate compromise: long enough to "
            "investigate an incident reported weeks after the fact, short enough that a "
            "log store never becomes an unmanaged copy of the production database. Teams "
            "that need longer must aggregate the data into a metric rather than extend "
            "the raw log retention.",
            "Support conversations, including chat transcripts and email threads, are "
            "retained for twenty-four months from the date the ticket is closed. "
            "Analytics events collected from the product are retained in raw form for "
            "ninety days and in aggregated form indefinitely, provided the aggregate "
            "cannot be resolved back to an individual user.",
        ],
    ),
    (
        "Deletion Procedure",
        [
            "Deletion is automated wherever the storage system supports it. Each system "
            "in scope declares its retention periods in configuration, and a scheduled "
            "job enforces them daily. Manual deletion is permitted only where automation "
            "is impossible, and every manual deletion is recorded in the change log with "
            "the operator, the scope and the reason.",
            "A deletion request received from a data subject is completed within twenty-one "
            "calendar days. The request is fanned out to every system holding the "
            "subject's data, and the requester receives a written confirmation listing "
            "which systems were in scope. A request that cannot be completed within the "
            "window is escalated to the data protection officer before the deadline "
            "passes, never after.",
        ],
    ),
    (
        "Backups and Restores",
        [
            "Backups are retained for thirty-five days on a rolling window. This is "
            "intentionally longer than most deletion windows, and it means a record "
            "deleted from production may still exist in a backup for up to five weeks "
            "afterwards. That gap is accepted, documented here, and disclosed in the "
            "privacy notice rather than quietly ignored.",
            "A restore from backup must not reintroduce deleted records. After any "
            "full-system restore, the deletion job is run immediately against the "
            "restored data set and the operator confirms in the incident record that it "
            "completed. A restore that skips this step is treated as a reportable "
            "incident in its own right.",
        ],
    ),
    (
        "Legal Hold and Exceptions",
        [
            "A legal hold suspends every retention rule for the data it names, including "
            "automated deletion, until the hold is lifted in writing by the legal "
            "function. A hold takes precedence over a data subject's deletion request; "
            "the requester is told that their data is subject to a hold, and the request "
            "is completed once the hold ends.",
            "Any other exception requires written approval from the data protection "
            "officer, a stated end date, and a named owner. Exceptions without an end "
            "date are not granted. The register of active exceptions is reviewed at every "
            "quarterly governance meeting and an exception whose owner has left the "
            "company is revoked automatically.",
        ],
    ),
    (
        "Ownership and Review",
        [
            "This policy is owned by the data protection officer and reviewed annually "
            "each January, or sooner if a change in law, a new processing activity or an "
            "incident makes the current text inaccurate. Proposed changes are circulated "
            "to engineering leads two weeks before they take effect.",
            "Each engineering team names one person accountable for retention compliance "
            "in the systems that team operates. That accountability sits with a named "
            "individual rather than with the team as a whole, because a control everyone "
            "owns is a control no one checks.",
        ],
    ),
]


def write_docx(path: Path) -> None:
    """Render the onboarding guide with real Heading 1/2 styles."""
    doc = Document()
    doc.add_heading("New Joiner Onboarding Guide", level=1)
    doc.add_paragraph(
        "This guide covers the first six months at the company: what happens before a "
        "start date, what the first week looks like, and how equipment, expenses and "
        "the probation review are handled."
    )
    for heading, paragraphs in ONBOARDING:
        doc.add_heading(heading, level=2)
        for para in paragraphs:
            doc.add_paragraph(para)
    doc.save(path)


def write_pdf(path: Path) -> None:
    """Render the retention policy as a multi-page, text-layer PDF."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:  # pragma: no cover - authoring-time only
        raise SystemExit(
            "reportlab is not installed (deliberately — nothing at runtime writes PDFs).\n"
            "Regenerate with: uv run --with reportlab python -m scripts.make_samples"
        ) from None

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        title="Data Retention Policy",
        author="DocTalk sample corpus",
    )
    flow = [Paragraph("Data Retention Policy", styles["Title"])]
    for heading, paragraphs in RETENTION:
        flow.append(Paragraph(heading, styles["Heading2"]))
        for para in paragraphs:
            flow.append(Paragraph(para, styles["BodyText"]))
        flow.append(Spacer(1, 8))
    doc.build(flow)


def main() -> None:
    docx_path = SAMPLES_DIR / "onboarding-guide.docx"
    pdf_path = SAMPLES_DIR / "data-retention-policy.pdf"
    write_docx(docx_path)
    write_pdf(pdf_path)
    for path in (docx_path, pdf_path):
        print(f"wrote {path.relative_to(SAMPLES_DIR.parent)} ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
