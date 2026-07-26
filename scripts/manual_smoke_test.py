"""Manual smoke test for DocTalk Phase 0-2 (ingestion + hybrid retrieval).

Exercises everything that exists at the code level today: no API/UI/LLM yet,
so this drives app.storage / app.retrieval directly, the same
way the future API layer eventually will. Read the output top to bottom.

Run from the project root:
    uv run python -m scripts.manual_smoke_test

Prerequisites:
    docker compose up -d db          # Postgres/ParadeDB reachable at DATABASE_URL

DESTRUCTIVE: step 0 drops and re-bootstraps the schema, so any documents you
uploaded manually into this database are lost. Point DATABASE_URL at a
throwaway database if you want to keep a demo workspace intact.
"""

import io
from pathlib import Path

import tiktoken
from docx import Document as DocxDocument

from app.chunking.chunker import _ENCODING
from app.config import settings
from app.retrieval import HybridRerankRetriever, embed_texts, embedding_dim
from app.retrieval.retriever import _dense_search, _lexical_search, _rrf_fuse
from app.storage import (
    WorkspaceFullError,
    count_documents,
    delete_document,
    get_chunks,
    get_connection,
    get_document,
    ingest_document,
    list_documents,
    reset_schema,
)

SEP = "=" * 78
ENC = tiktoken.get_encoding(_ENCODING)


def section(title: str) -> None:
    print(f"\n{SEP}\n{title}\n{SEP}")


def ok(msg: str) -> None:
    print(f"  OK   - {msg}")


def note(msg: str) -> None:
    print(f"  NOTE - {msg}")


def tokens(text: str) -> int:
    return len(ENC.encode(text))


def main() -> None:
    # ---------------------------------------------------------------------------
    section("0. Reset schema by replaying migrations/ (scripts/reset_db.py)")
    versions = reset_schema()
    ok(
        f"schema (re)created from {', '.join(versions)}; "
        f"embedding_dim={embedding_dim()} ({settings.embedding_model}), "
        f"max_documents={settings.max_documents}"
    )

    # ---------------------------------------------------------------------------
    section("1. Ingest a Markdown document")
    md_content = b"""# Sick Leave Policy

## Eligibility

All full-time employees are eligible for sick leave starting on day one of
employment. Part-time employees accrue leave on a pro-rated basis.

## Vacation

Full-time employees accrue fifteen days of vacation per year. Vacation
requests must be submitted two weeks in advance to a manager.
"""
    result_md = ingest_document("hr-policy.md", md_content)
    print(f"  {result_md}")
    assert result_md.chunk_count > 0 and not result_md.deduplicated
    ok("MD ingested, chunked, embedded, stored")

    # ---------------------------------------------------------------------------
    section("2. Ingest a DOCX document (built in-memory with python-docx)")
    docx_buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Password Policy", level=1)
    doc.add_heading("Requirements", level=2)
    doc.add_paragraph(
        "Passwords must be at least twelve characters and rotated every ninety "
        "days. Multi-factor authentication is required for all remote access."
    )
    doc.add_heading("VPN Access", level=2)
    doc.add_paragraph(
        "Employees connecting from outside the office must use the corporate "
        "VPN client, configured by the IT helpdesk."
    )
    doc.save(docx_buf)
    result_docx = ingest_document("it-policy.docx", docx_buf.getvalue())
    print(f"  {result_docx}")
    assert result_docx.chunk_count > 0
    ok("DOCX ingested")

    # ---------------------------------------------------------------------------
    section("3. Ingest a real PDF (the assignment brief already in docs/)")
    pdf_path = Path("docs/Case Study Assignment - AI Engineer.pdf")
    result_pdf = ingest_document(pdf_path.name, pdf_path.read_bytes())
    print(f"  {result_pdf}")
    assert result_pdf.chunk_count > 0
    ok("PDF ingested")

    # ---------------------------------------------------------------------------
    section("4. Re-upload identical MD bytes -> expect a dedup no-op, no new rows")
    result_dup = ingest_document("hr-policy.md", md_content)
    print(f"  {result_dup}")
    assert result_dup.deduplicated and result_dup.document_id == result_md.document_id
    assert result_dup.chunk_count == result_md.chunk_count
    ok("duplicate upload recognized, no new document/chunks created")

    # ---------------------------------------------------------------------------
    section("4b. OBSERVATION: same bytes, different filename -> NOT deduped")
    result_relabel = ingest_document("hr-policy-v2.md", md_content)
    print(f"  {result_relabel}")
    assert result_relabel.document_id != result_md.document_id
    note(
        "make_document_id() mixes the filename stem into the id, so dedup is keyed on "
        "(filename, content), not content alone: the same file renamed burns a second "
        "workspace slot. Documented as an assumption in the README."
    )
    with get_connection() as conn:
        delete_document(conn, result_relabel.document_id)
    note("deleted the renamed copy again so it does not pollute the retrieval steps")

    # ---------------------------------------------------------------------------
    section("5. Reject an unsupported file type")
    try:
        ingest_document("notes.txt", b"whatever")
        print("  FAIL: expected ValueError, none was raised")
    except ValueError as e:
        ok(f"rejected as expected: {e}")

    # ---------------------------------------------------------------------------
    section("6. Reject a document with no extractable text")
    try:
        ingest_document("empty.md", b"   \n\n   ")
        print("  FAIL: expected ValueError, none was raised")
    except ValueError as e:
        ok(f"rejected as expected: {e}")

    # ---------------------------------------------------------------------------
    section("7. Inspect workspace state (documents table)")
    with get_connection() as conn:
        for d in list_documents(conn):
            print(f"  {d}")
        print(f"  Total documents: {count_documents(conn)}")

    # ---------------------------------------------------------------------------
    section("8. Inspect chunk structure per format (IDs, section_path, page, tokens)")
    with get_connection() as conn:
        md_chunks = get_chunks(conn, result_md.document_id)
        docx_chunks = get_chunks(conn, result_docx.document_id)
        pdf_chunks = get_chunks(conn, result_pdf.document_id)

    print("\n  --- MD: expect a heading-derived section_path, page_number None ---")
    for c in md_chunks:
        print(
            f"  {c['id']:34s} tokens={tokens(c['text']):4d} section={c['section_path']} "
            f"page={c['page_number']} chars={c['char_start']}-{c['char_end']}"
        )
        print(f"      {c['text'][:70]!r}")
    assert all(c["section_path"] for c in md_chunks), "expected MD chunks to carry a heading path"
    assert all(c["page_number"] is None for c in md_chunks)

    print("\n  --- DOCX: expect a Heading-style section_path, page_number None ---")
    for c in docx_chunks:
        print(f"  {c['id']:34s} tokens={tokens(c['text']):4d} section={c['section_path']}")
        print(f"      {c['text'][:70]!r}")
    assert all(c["section_path"] for c in docx_chunks), (
        "expected DOCX chunks to carry a heading path"
    )

    print("\n  --- PDF: expect page_number set, section_path None (no heading concept) ---")
    for c in pdf_chunks[:4]:
        print(
            f"  {c['id']:34s} tokens={tokens(c['text']):4d} page={c['page_number']} "
            f"section={c['section_path']}"
        )
        print(f"      head={c['text'][:60]!r}")
        print(f"      tail={c['text'][-60:]!r}")
    assert all(c["page_number"] is not None for c in pdf_chunks), (
        "expected every PDF chunk to carry a page"
    )
    assert all(tokens(c["text"]) <= 256 for c in pdf_chunks), (
        "chunks must fit the embedding model's 256-token window, or their tails are "
        "silently truncated at embed time"
    )
    note(
        "eyeball the tail of one PDF chunk reappearing at the head of the next: that is "
        "the 12.5% sliding overlap (~31 tokens)"
    )
    ok("chunk IDs, section paths, page numbers and token budget all as expected")

    # ---------------------------------------------------------------------------
    section("9. Verify what actually landed in Postgres (vectors + indexes)")
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT DISTINCT vector_dims(embedding) FROM chunks")
            dims = [r[0] for r in cur.fetchall()]
            cur.execute("SELECT count(*), count(DISTINCT embedding::text) FROM chunks")
            total, distinct = cur.fetchone()
            cur.execute("SELECT count(*) FROM chunks WHERE embedding IS NULL")
            nulls = cur.fetchone()[0]
            cur.execute("SELECT indexname FROM pg_indexes WHERE tablename = 'chunks'")
            indexes = sorted(r[0] for r in cur.fetchall())

            # A chunk must be its own nearest neighbour at ~zero cosine distance.
            cur.execute(
                """
            WITH probe AS (SELECT id, embedding FROM chunks ORDER BY id LIMIT 1)
            SELECT c.id, round((c.embedding <=> p.embedding)::numeric, 4)
            FROM chunks c, probe p ORDER BY c.embedding <=> p.embedding LIMIT 3
            """
            )
            neighbours = cur.fetchall()

    print(f"  vector_dims present: {dims}")
    print(f"  chunk rows: {total}, distinct embeddings: {distinct}, NULL embeddings: {nulls}")
    print(f"  indexes on chunks: {indexes}")
    print("  nearest neighbours of the first chunk (self must be #1 at ~0.0):")
    for cid, dist in neighbours:
        print(f"    {dist}  {cid}")

    assert dims == [embedding_dim()], f"expected only {embedding_dim()}-dim vectors"
    assert nulls == 0
    assert distinct == total, "identical embeddings across chunks means the embed step is broken"
    assert "chunks_embedding_idx" in indexes and "chunks_bm25_idx" in indexes
    assert float(neighbours[0][1]) < 0.001, "a chunk is not its own nearest neighbour"
    ok("embeddings are right-sized, non-null, distinct; HNSW + BM25 indexes exist")

    # Norm via the negative-inner-product operator: |v| = sqrt(-(v <#> v)).
    # `l2_norm(vector)` is ambiguous in this pgvector/ParadeDB build.
    with get_connection() as conn, conn.cursor() as cur:
        cur.execute("SELECT round(sqrt(-(embedding <#> embedding))::numeric, 4) FROM chunks")
        norms = [float(r[0]) for r in cur.fetchall()]
    print(f"  L2 norms (expect all ~1.0, embeddings are normalized): {sorted(set(norms))}")
    assert all(abs(n - 1.0) < 0.01 for n in norms), "embeddings are not L2-normalized"
    ok("embeddings are L2-normalized, so cosine distance agrees with dot product")

    # ---------------------------------------------------------------------------
    section(f"10. Enforce the workspace cap (max_documents = {settings.max_documents})")
    filler_ids: list[str] = []
    try:
        for i in range(10):
            content = f"# Filler Doc {i}\n\nThis is filler content number {i}.".encode()
            filler_ids.append(ingest_document(f"filler-{i}.md", content).document_id)
            print(f"  ingested filler-{i}.md -> {filler_ids[-1]}")
    except WorkspaceFullError as e:
        ok(f"cap enforced as expected: {e}")

    with get_connection() as conn:
        assert count_documents(conn) == settings.max_documents
    print(f"  Workspace now at {settings.max_documents}/{settings.max_documents} documents")
    assert filler_ids, "expected at least one filler to fit before the cap was reached"

    # ---------------------------------------------------------------------------
    section("10b. Delete one document, confirm cascade + freed slot")
    victim_id = filler_ids[0]
    with get_connection() as conn:
        was_deleted = delete_document(conn, victim_id)
        remaining_chunks = get_chunks(conn, victim_id)
        doc_gone = get_document(conn, victim_id) is None
        count_after_delete = count_documents(conn)
    print(
        f"  delete_document -> {was_deleted}, chunks left -> {len(remaining_chunks)}, "
        f"doc row gone -> {doc_gone}, count now -> {count_after_delete}"
    )
    assert was_deleted and not remaining_chunks and doc_gone
    assert count_after_delete == settings.max_documents - 1
    ok("delete cascades to chunks and frees a workspace slot")

    reslot = ingest_document("filler-new.md", b"# Filler New\n\nFresh filler content.")
    print(f"  re-ingested into freed slot -> {reslot}")
    ok("freed slot can be reused")

    # ---------------------------------------------------------------------------
    section("11. Retrieval, leg by leg (dense -> lexical -> RRF -> rerank)")
    retriever = HybridRerankRetriever()
    for query in [
        "How many vacation days do full-time employees get?",
        "What are the password requirements for remote access?",
    ]:
        print(f"\n  --- Query: {query!r} ---")
        query_vector = embed_texts([query])[0]
        with get_connection() as conn:
            dense_ids = _dense_search(conn, query_vector, 10)
            lexical_ids = _lexical_search(conn, query, 10)
            fused_ids = _rrf_fuse([dense_ids, lexical_ids])[:10]
            final = retriever.retrieve(conn, query, top_k=3)
        print(f"    dense   ({len(dense_ids):2d}): {dense_ids[:5]}")
        print(f"    lexical ({len(lexical_ids):2d}): {lexical_ids[:5]}")
        print(f"    fused   ({len(fused_ids):2d}): {fused_ids[:5]}")
        print("    reranked top-3:")
        for r in final:
            print(
                f"      [{r.rerank_score:+.3f}] dense_rank={r.dense_rank} "
                f"lexical_rank={r.lexical_rank}  {r.chunk_id} ({r.filename})"
            )
            print(f"          {r.text[:90]!r}")
        assert final, f"expected results for {query!r}"
    note("compare the four orderings: fusion should mix both legs, rerank should reshuffle them")

    # ---------------------------------------------------------------------------
    section("12. Prove the lexical leg is alive, not silently returning nothing")
    with get_connection() as conn:
        lexical_only = _lexical_search(conn, "helpdesk", 10)
        hybrid = retriever.retrieve(conn, "helpdesk", top_k=3)
    print(f"  BM25 hits for 'helpdesk': {lexical_only}")
    assert lexical_only, (
        "pg_search/BM25 returned nothing for an exact rare term - lexical leg is dead"
    )
    assert any(r.lexical_rank is not None for r in hybrid), (
        "lexical leg contributed nothing to the fusion"
    )
    ok("pg_search/BM25 leg returns hits and reaches the fused result")

    # ---------------------------------------------------------------------------
    section("13. Why hybrid: a paraphrase with no lexical overlap")
    with get_connection() as conn:
        paraphrase = retriever.retrieve(conn, "annual time off allowance", top_k=3)
    for r in paraphrase:
        print(
            f"    [{r.rerank_score:+.3f}] dense_rank={r.dense_rank} "
            f"lexical_rank={r.lexical_rank}  {r.chunk_id} ({r.filename})"
        )
    assert paraphrase[0].filename == "hr-policy.md", (
        f"expected the vacation chunk on top, got {paraphrase[0].chunk_id}"
    )
    ok("dense leg carries a query that shares no keywords with the source text")

    # ---------------------------------------------------------------------------
    section("14. Adversarial punctuation must not reach ParadeDB's query parser")
    raised: list[str] = []
    for query in ["vacation policy: 15 days?", "-vacation", "C++", 'unclosed "quote', "a AND (b"]:
        try:
            with get_connection() as conn:
                hits = retriever.retrieve(conn, query, top_k=1)
            print(f"    {query!r} -> {len(hits)} hit(s)")
        except Exception as e:  # noqa: BLE001 - collected and asserted below
            print(f"    RAISED {query!r} -> {type(e).__name__}: {str(e).splitlines()[0]}")
            raised.append(query)
    assert not raised, f"lexical leg still parses user punctuation as query syntax: {raised}"
    ok("punctuated queries survive: `paradedb.match` takes the input as plain terms")

    # ---------------------------------------------------------------------------
    section("15. Off-topic query should score lower than an on-topic one")
    with get_connection() as conn:
        on_topic = retriever.retrieve(conn, "vacation days for employees", top_k=1)
        off_topic = retriever.retrieve(conn, "what is the airspeed velocity of a swallow", top_k=1)
    print(f"  on-topic  {on_topic[0].rerank_score:+.3f}  ({on_topic[0].chunk_id})")
    print(f"  off-topic {off_topic[0].rerank_score:+.3f}  ({off_topic[0].chunk_id})")
    assert on_topic[0].rerank_score > off_topic[0].rerank_score
    note(
        "dense search always returns its top-k, so an off-topic query still yields rows - "
        "the rerank score, not the presence of results, is what Phase 3 must gate on"
    )
    ok("reranker separates on-topic from off-topic as expected")

    # ---------------------------------------------------------------------------
    section("16. Deleted content must disappear from retrieval (no stale vectors)")
    with get_connection() as conn:
        delete_document(conn, result_docx.document_id)
        after_delete = retriever.retrieve(
            conn, "passwords twelve characters remote access", top_k=5
        )
    for r in after_delete:
        print(f"    [{r.rerank_score:+.3f}] {r.chunk_id} ({r.filename})")
    assert all(r.document_id != result_docx.document_id for r in after_delete), (
        "a deleted document is still retrievable - stale vectors would let the "
        "LLM cite removed content"
    )
    ok("deleted document's chunks are gone from the index and from retrieval")

    section("DONE. Everything above should read OK / NOTE, with no FAIL or traceback.")


if __name__ == "__main__":
    main()
