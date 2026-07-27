"""Parsing PDF/DOCX/MD into blocks — the first step of the citation chain.

Structure extracted here becomes the `section_path` and `page_number` a reader
sees on a citation chip, and the char offsets are what let a chunk be located
in the original file. A parser that quietly drops structure produces citations
that still validate but point nowhere useful, so these assert the structure,
not just that some text came out.
"""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pdfplumber.utils.exceptions import PdfminerException

from app.parsers import parse_document
from app.parsers.markdown import parse_markdown

PDF_FIXTURE = (
    Path(__file__).resolve().parents[2] / "docs" / "Case Study Assignment - AI Engineer.pdf"
)

MARKDOWN = """\
# HR Policy

Introductory sentence.

## Vacation

Full-time employees accrue fifteen days per year.
Requests go to a line manager.

### Carry-over

Capped at five days.

## Sick Leave

A doctor's note is required from day four.
"""


def section_paths(blocks):
    return [block.section_path for block in blocks]


# --------------------------------------------------------------------- markdown


def test_markdown_nests_headings_into_a_section_path():
    blocks = parse_markdown("hr.md", MARKDOWN).blocks

    assert section_paths(blocks) == [
        ["HR Policy"],
        ["HR Policy", "Vacation"],
        ["HR Policy", "Vacation", "Carry-over"],
        ["HR Policy", "Sick Leave"],
    ]


def test_a_shallower_heading_pops_the_deeper_ones():
    """`## Sick Leave` must not stay nested under `### Carry-over`."""
    blocks = parse_markdown("hr.md", MARKDOWN).blocks

    assert blocks[-1].section_path == ["HR Policy", "Sick Leave"]


def test_consecutive_lines_stay_one_block_and_blank_lines_split_them():
    blocks = parse_markdown("hr.md", MARKDOWN).blocks

    vacation = blocks[1]
    assert vacation.text == (
        "Full-time employees accrue fifteen days per year.\nRequests go to a line manager."
    )
    assert len(blocks) == 4


def test_char_offsets_locate_the_block_in_the_original_source():
    """The offsets are how a chunk is traced back to the uploaded file."""
    blocks = parse_markdown("hr.md", MARKDOWN).blocks

    for block in blocks:
        assert MARKDOWN[block.char_start : block.char_end].strip() == block.text


def test_headings_are_structure_not_content():
    """A heading names the section it labels; it is not a block of its own."""
    blocks = parse_markdown("hr.md", "# Only A Heading\n").blocks

    assert blocks == []


def test_markdown_without_any_heading_still_parses():
    blocks = parse_markdown("notes.md", "Just a loose note.\n").blocks

    assert len(blocks) == 1
    assert blocks[0].section_path is None
    assert blocks[0].page_number is None


# ------------------------------------------------------------------------ docx


def build_docx(items: list[tuple[str, str]]) -> bytes:
    """`items` are (style, text); style "" means a normal paragraph."""
    document = Document()
    for style, text in items:
        if style:
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_heading_styles_build_the_section_path():
    content = build_docx(
        [
            ("Heading 1", "HR Policy"),
            ("Heading 2", "Vacation"),
            ("", "Fifteen days per year."),
            ("Heading 2", "Sick Leave"),
            ("", "A note is required from day four."),
        ]
    )

    blocks = parse_document("hr.docx", content).blocks

    assert [b.text for b in blocks] == [
        "Fifteen days per year.",
        "A note is required from day four.",
    ]
    assert section_paths(blocks) == [
        ["HR Policy", "Vacation"],
        ["HR Policy", "Sick Leave"],
    ]


def test_docx_skips_empty_paragraphs():
    content = build_docx([("", "Real content."), ("", "   "), ("", "More content.")])

    assert [b.text for b in parse_document("x.docx", content).blocks] == [
        "Real content.",
        "More content.",
    ]


def test_docx_carries_no_page_numbers():
    """Pagination is a rendering decision, not present in the DOCX source."""
    blocks = parse_document("x.docx", build_docx([("", "Body.")])).blocks

    assert blocks[0].page_number is None


# ------------------------------------------------------------------------- pdf


@pytest.mark.skipif(not PDF_FIXTURE.exists(), reason="PDF fixture not present")
def test_pdf_extracts_text_and_tags_every_block_with_its_page():
    blocks = parse_document(PDF_FIXTURE.name, PDF_FIXTURE.read_bytes()).blocks

    assert blocks, "no text extracted from a PDF that has a text layer"
    assert all(b.page_number is not None for b in blocks)
    # 1-based, and non-decreasing: a citation's "(p. 4)" has to mean page 4.
    assert min(b.page_number for b in blocks) == 1
    pages = [b.page_number for b in blocks]
    assert pages == sorted(pages)
    assert all(b.text.strip() for b in blocks)


def test_a_corrupt_pdf_raises_rather_than_yielding_an_empty_document():
    """A file that only claims to be a PDF must fail, not ingest as empty.

    The neighbouring case — a *valid* PDF with no text layer, i.e. a scan —
    parses to zero blocks and is rejected one layer up, at ingestion; that
    path is covered in `tests/integration/test_api.py`.
    """
    with pytest.raises(PdfminerException):
        parse_document("scan.pdf", b"%PDF-1.4 not really a pdf")


# -------------------------------------------------------------------- dispatch


def test_the_extension_selects_the_parser_case_insensitively():
    content = build_docx([("", "Body.")])

    assert parse_document("UPPER.DOCX", content).blocks[0].text == "Body."
    assert parse_document("Notes.MD", b"# T\n\nBody.\n").blocks[0].text == "Body."


@pytest.mark.parametrize("filename", ["notes.txt", "archive.zip", "noextension"])
def test_unsupported_types_are_rejected_by_name(filename):
    with pytest.raises(ValueError, match="Supported: PDF, DOCX, MD"):
        parse_document(filename, b"whatever")


def test_undecodable_markdown_bytes_raise_rather_than_mangle_text():
    """The upload route turns this into a per-file rejection, not a 500."""
    with pytest.raises(UnicodeDecodeError):
        parse_document("bad.md", b"\xff\xfe\x00invalid")


def test_the_parsed_document_keeps_the_filename_it_was_given():
    assert parse_document("report.md", b"# T\n\nBody.\n").filename == "report.md"
