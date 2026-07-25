import re
from io import BytesIO

from docx import Document

from .base import Block, ParsedDocument

_HEADING_RE = re.compile(r"^Heading (\d)$")


def parse_docx(filename: str, content: bytes) -> ParsedDocument:
    doc = Document(BytesIO(content))
    blocks: list[Block] = []
    section_stack: list[str] = []
    offset = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        heading_match = _HEADING_RE.match(style_name)
        start, end = offset, offset + len(text)
        if heading_match:
            level = int(heading_match.group(1))
            section_stack = section_stack[: level - 1]
            section_stack.append(text)
        else:
            blocks.append(
                Block(
                    text=text,
                    section_path=list(section_stack) or None,
                    page_number=None,
                    char_start=start,
                    char_end=end,
                )
            )
        offset = end + 2  # matches the "\n\n" join used for offset bookkeeping

    return ParsedDocument(filename=filename, blocks=blocks)
